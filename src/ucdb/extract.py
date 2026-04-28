"""Extract plain text from PDF, DOCX, ODT, plain-text, and Markdown documents."""

from __future__ import annotations

from pathlib import Path

PLAINTEXT_EXTENSIONS = {".txt", ".text", ".md", ".markdown", ".mdown", ".mkd"}
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".odt"} | PLAINTEXT_EXTENSIONS


class UnsupportedFormatError(ValueError):
    pass


def extract_text(path: Path | str) -> str:
    """Return the plain-text content of a supported document.

    Markdown is passed through verbatim — the AI step preserves structure
    better when it sees the original headings and list markers.
    """
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".odt":
        return _extract_odt(path)
    if suffix in PLAINTEXT_EXTENSIONS:
        return path.read_text(encoding="utf-8", errors="replace")
    raise UnsupportedFormatError(
        f"Unsupported document format: {suffix} (supported: {sorted(SUPPORTED_EXTENSIONS)})"
    )


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            parts.append("\t".join(cells))
    return "\n".join(parts)


def _extract_odt(path: Path) -> str:
    from odf import teletype
    from odf import text as odftext
    from odf.opendocument import load

    doc = load(str(path))
    parts: list[str] = []
    for elem in doc.getElementsByType(odftext.P):
        parts.append(teletype.extractText(elem))
    for elem in doc.getElementsByType(odftext.H):
        parts.append(teletype.extractText(elem))
    return "\n".join(p for p in parts if p)
